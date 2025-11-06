"""
AHKStrategies ULTIMATE Letterhead Generator
HTML→PDF (with animations), plus clean DOCX with proper footer
BEST OF BOTH WORLDS
"""

import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_top_border(paragraph, color='D4AF37', size=12):
    """Add a top border to a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:space'), '8')
    top.set(qn('w:color'), color)
    pBdr.append(top)
    pPr.append(pBdr)

def create_docx_with_footer():
    """Create DOCX with styled header and proper Word footer"""
    
    # Create new document
    doc = Document()
    
    # Set page margins (A4: 210mm x 297mm)
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    
    # ===== HEADER SECTION (in document body) =====
    # Company Name
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run('AHKSTRATEGIES')
    run.font.name = 'Playfair Display'
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(212, 175, 55)
    run.font.bold = True
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(2)
    
    # Tagline
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.add_run('Where Human Brilliance Fuses with AI Symphony')
    run.font.name = 'Cormorant Garamond'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(160, 130, 109)
    run.font.italic = True
    h2.paragraph_format.space_after = Pt(2)
    
    # Vision Statement
    h3 = doc.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h3.add_run('BUILDING EMPIRES • CRAFTING FUTURES • TRANSCENDING LIMITS')
    run.font.name = 'Inter'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(139, 115, 85)
    h3.paragraph_format.space_after = Pt(16)
    
    # Golden separator line
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = separator.add_run('━' * 80)
    run.font.color.rgb = RGBColor(212, 175, 55)
    run.font.size = Pt(6)
    separator.paragraph_format.space_after = Pt(24)
    
    # Content area placeholder (user can type here)
    content = doc.add_paragraph()
    content.paragraph_format.space_after = Pt(300)
    
    # ===== FOOTER SECTION (as proper Word footer) =====
    footer = section.footer
    
    # Line 1 with golden top border
    p1 = footer.paragraphs[0]
    add_top_border(p1, 'D4AF37', 12)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(4)
    
    # Company name
    run = p1.add_run('AHKStrategies')
    run.font.name = 'Playfair Display'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(212, 175, 55)
    
    run = p1.add_run(' ◆ ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(212, 175, 55)
    
    # Address
    run = p1.add_run('Cairo • Dubai • Amman')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    run = p1.add_run(' • ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Phone
    run = p1.add_run('+20 104 078 7571')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    run = p1.add_run(' • ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Website
    run = p1.add_run('www.ahkstrategies.net')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(160, 130, 109)
    run.font.italic = True
    
    # Line 2: Copyright
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    
    run = p2.add_run('Confidential & Proprietary © 2025 AHKStrategies – All Rights Reserved')
    run.font.name = 'Inter'
    run.font.size = Pt(6.5)
    run.font.color.rgb = RGBColor(153, 153, 153)
    
    # Save
    docx_path = 'AHKStrategies_Letterhead_LEGENDARY.docx'
    doc.save(docx_path)
    return docx_path

async def html_to_pdf_to_docx():
    """Generate both PDF (animated) and DOCX (editable) versions"""
    
    html_path = r"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\letterhead_legendary.html"
    pdf_path = r"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\AHKStrategies_Letterhead_LEGENDARY.pdf"
    
    print("=" * 100)
    print("🔥 LEGENDARY LETTERHEAD GENERATION - PERFECTION IN TWO FORMATS")
    print("=" * 100)
    print("📄 Step 1: Rendering HTML with Chromium (cinema-quality with animations)...")
    
    async with async_playwright() as p:
        # Launch browser
        browser = await p.chromium.launch()
        page = await browser.new_page()
        
        # Load HTML file
        await page.goto(f'file:///{html_path}')
        
        # Wait for fonts and animations to load
        await page.wait_for_timeout(3000)
        
        # Generate PDF with maximum quality
        await page.pdf(
            path=pdf_path,
            format='A4',
            print_background=True,
            prefer_css_page_size=True,
            margin={'top': '0mm', 'right': '0mm', 'bottom': '0mm', 'left': '0mm'}
        )
        
        await browser.close()
    
    print(f"✅ Step 1 Complete: Animated PDF with quantum neural design")
    print(f"📄 Step 2: Creating clean editable DOCX with proper footer...")
    
    # Create DOCX with proper footer
    docx_path = create_docx_with_footer()
    
    print(f"✅ Step 2 Complete: DOCX with editable content area and Word footer")
    print("\n" + "=" * 100)
    print("✨✨✨ LEGENDARY LETTERHEAD COMPLETE ✨✨✨")
    print("=" * 100)
    print(f"📄 HTML: {html_path}")
    print(f"📄 PDF:  {pdf_path}")
    print(f"   ↳ WITH QUANTUM RINGS, NEURAL PATHWAYS, ANIMATIONS")
    print(f"📄 DOCX: {os.path.abspath(docx_path)}")
    print(f"   ↳ CLEAN EDITABLE FORMAT WITH PROPER WORD FOOTER")
    print("=" * 100)
    print("🏆 PDF: View/Print with full animated quantum neural design")
    print("💎 DOCX: Edit and type your content - footer stays at bottom")
    print("🔥 STATUS: READY TO DOMINATE THE UNIVERSE")
    print("=" * 100)

if __name__ == "__main__":
    asyncio.run(html_to_pdf_to_docx())
