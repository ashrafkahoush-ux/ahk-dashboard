"""
Create Clean DOCX Letterhead from Scratch
With proper header, editable content area, and Word footer
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_top_border(paragraph, color='D4AF37', size=12):
    """Add a top border to a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:space'), '8')
    top.set(qn('w:color'), color)
    pBdr.append(top)
    pPr.append(pBdr)

def create_letterhead():
    """Create the letterhead with header and footer"""
    
    # Create new document
    doc = Document()
    
    # Set page margins (A4: 210mm x 297mm)
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    
    # ===== HEADER SECTION (in document body, not Word header) =====
    # Company Name
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run('AHKSTRATEGIES')
    run.font.name = 'Playfair Display'
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(212, 175, 55)  # Gold
    run.font.bold = True
    h1.paragraph_format.space_after = Pt(2)
    
    # Tagline
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.add_run('Where Human Brilliance Fuses with AI Symphony')
    run.font.name = 'Cormorant Garamond'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(160, 130, 109)
    run.font.italic = True
    h2.paragraph_format.space_after = Pt(2)
    
    # Vision Statement
    h3 = doc.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h3.add_run('BUILDING EMPIRES • CRAFTING FUTURES • TRANSCENDING LIMITS')
    run.font.name = 'Inter'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(139, 115, 85)
    h3.paragraph_format.space_after = Pt(24)
    
    # Add some spacing for content area
    content = doc.add_paragraph()
    content.paragraph_format.space_after = Pt(300)  # Large space for content
    
    # ===== FOOTER SECTION (as Word footer) =====
    footer = section.footer
    
    # Line 1 with golden top border
    p1 = footer.paragraphs[0]
    add_top_border(p1, 'D4AF37', 12)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(4)
    
    # Company name
    run = p1.add_run('AHKStrategies')
    run.font.name = 'Playfair Display'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(212, 175, 55)
    
    run = p1.add_run(' ◆ ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(212, 175, 55)
    
    # Address
    run = p1.add_run('Cairo • Dubai • Amman')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    run = p1.add_run(' • ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Phone
    run = p1.add_run('+20 104 078 7571')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    run = p1.add_run(' • ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Website
    run = p1.add_run('www.ahkstrategies.net')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(160, 130, 109)
    run.font.italic = True
    
    # Line 2: Copyright
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    
    run = p2.add_run('Confidential & Proprietary © 2025 AHKStrategies – All Rights Reserved')
    run.font.name = 'Inter'
    run.font.size = Pt(6.5)
    run.font.color.rgb = RGBColor(153, 153, 153)
    
    # Save
    doc.save('AHKStrategies_Letterhead_LEGENDARY.docx')
    print("✅ Clean letterhead created!")
    print("📄 File: AHKStrategies_Letterhead_LEGENDARY.docx")
    print("✨ You can now type your content in the middle section")

if __name__ == "__main__":
    create_letterhead()
