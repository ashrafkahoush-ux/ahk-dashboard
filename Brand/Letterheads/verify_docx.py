from docx import Document

doc = Document('AHKStrategies_Letterhead_LEGENDARY.docx')

print('='*60)
print('DOCX VERIFICATION')
print('='*60)
print(f'Sections: {len(doc.sections)}')
print(f'Body paragraphs: {len(doc.paragraphs)}')

print('\nBody content:')
for i in range(min(6, len(doc.paragraphs))):
    text = doc.paragraphs[i].text
    if text:
        print(f'  Para {i}: {text[:80]}...' if len(text) > 80 else f'  Para {i}: {text}')
    else:
        print(f'  Para {i}: [empty paragraph for content]')

footer = doc.sections[0].footer
print(f'\nFooter paragraphs: {len(footer.paragraphs)}')
print('Footer content:')
for i, p in enumerate(footer.paragraphs):
    text = p.text
    if text:
        print(f'  Line {i+1}: {text}')
    else:
        print(f'  Line {i+1}: [empty]')

print('='*60)
print('✅ VERIFICATION COMPLETE')
print('='*60)
