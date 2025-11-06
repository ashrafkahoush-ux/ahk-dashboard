"""
Add Footer to DOCX Letterhead
Adds the two-line footer directly to the Word document
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_footer_border(section):
    """Add a golden border line above the footer"""
    # Get the footer
    footer = section.footer
    
    # Add border to the first paragraph
    if len(footer.paragraphs) == 0:
        footer.add_paragraph()
    
    first_para = footer.paragraphs[0]
    pPr = first_para._element.get_or_add_pPr()
    
    # Create border element
    pBdr = OxmlElement('w:pBdr')
    
    # Top border (golden line)
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '12')  # 1.5pt (12/8)
    top_border.set(qn('w:space'), '8')  # Space below border
    top_border.set(qn('w:color'), 'D4AF37')  # Gold color
    
    pBdr.append(top_border)
    pPr.append(pBdr)

def add_footer_to_letterhead():
    """Add professional footer to the letterhead DOCX"""
    
    # Open the document
    doc = Document('AHKStrategies_Letterhead_LEGENDARY.docx')
    
    # Access the default section
    section = doc.sections[0]
    
    # Add footer border
    add_footer_border(section)
    
    # Get or create footer
    footer = section.footer
    
    # Clear existing footer paragraphs except the first (which has the border)
    while len(footer.paragraphs) > 1:
        footer._element.remove(footer.paragraphs[-1]._element)
    
    # Line 1: Company name, address, phone, website
    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company name
    run = p1.add_run('AHKStrategies')
    run.font.name = 'Playfair Display'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(212, 175, 55)  # Gold
    run.font.bold = False
    
    # Separator
    run = p1.add_run(' ◆ ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(212, 175, 55)
    
    # Address icon and text
    run = p1.add_run('📍 Cairo • Dubai • Amman')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Separator
    run = p1.add_run(' • ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Phone icon and number
    run = p1.add_run('☎ +20 104 078 7571')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Separator
    run = p1.add_run(' • ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    # Website icon and URL
    run = p1.add_run('🌐 ')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    
    run = p1.add_run('www.ahkstrategies.net')
    run.font.name = 'Inter'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(160, 130, 109)  # Muted gold
    run.font.italic = True
    
    # Line 2: Copyright
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p2.add_run('Confidential & Proprietary © 2025 AHKStrategies – All Rights Reserved')
    run.font.name = 'Inter'
    run.font.size = Pt(6.5)
    run.font.color.rgb = RGBColor(153, 153, 153)  # Gray
    
    # Adjust spacing
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    
    # Save the document
    doc.save('AHKStrategies_Letterhead_LEGENDARY.docx')
    print("✅ Footer added successfully to DOCX!")
    print("📄 File: AHKStrategies_Letterhead_LEGENDARY.docx")

if __name__ == "__main__":
    add_footer_to_letterhead()
