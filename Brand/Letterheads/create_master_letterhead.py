"""
AHKStrategies Master Letterhead Generator
The ULTIMATE elegant, sophisticated corporate letterhead with premium typography
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os

def add_geometric_line(paragraph, color_hex="#d4b37f", width_pt=1.5, style='single'):
    """Add a sophisticated line to a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), style)
    bottom.set(qn('w:sz'), str(int(width_pt * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex.strip('#'))
    
    pBdr.append(bottom)
    pPr.append(pBdr)

def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def add_character_spacing(run, spacing):
    """Add character spacing (letter spacing) to a run - spacing in twips (1/20 pt)"""
    rPr = run._element.get_or_add_rPr()
    spacing_elem = OxmlElement('w:spacing')
    spacing_elem.set(qn('w:val'), str(spacing))
    rPr.append(spacing_elem)

def create_master_letterhead():
    """Create the ULTIMATE AHKStrategies Master Letterhead"""
    
    # Initialize document
    doc = Document()
    
    # Set page margins - A4 with 1.8cm margins
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)  # A4 height
        section.page_width = Cm(21)     # A4 width
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    
    # === PREMIUM HEADER SECTION ===
    
    # Minimal breathing space at top
    top_space = doc.add_paragraph()
    top_space.paragraph_format.space_after = Pt(8)
    
    # Company Name - AHKSTRATEGIES properly on one line
    company_name = doc.add_paragraph()
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_name.paragraph_format.space_before = Pt(0)
    company_name.paragraph_format.space_after = Pt(4)
    
    name_run = company_name.add_run('A H K S T R A T E G I E S')
    name_run.font.name = 'Cormorant Garamond'
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(*hex_to_rgb('#111111'))
    name_run.font.bold = False
    
    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_before = Pt(0)
    tagline.paragraph_format.space_after = Pt(8)
    
    tagline_run = tagline.add_run('Strategic   •   Human   •   Intelligent')
    tagline_run.font.name = 'Montserrat'
    tagline_run.font.size = Pt(8)
    tagline_run.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    
    # Add the CROPPED banner image - proper quality with transparent background
    banner_para = doc.add_paragraph()
    banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_para.paragraph_format.space_before = Pt(0)
    banner_para.paragraph_format.space_after = Pt(14)
    
    banner_path = r"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\banner_transparent.png"
    if os.path.exists(banner_path):
        run = banner_para.add_run()
        run.add_picture(banner_path, width=Cm(14))  # Proper width for cropped banner
    
    # Elegant thin separator after banner
    separator = doc.add_paragraph()
    add_geometric_line(separator, color_hex="#d4b37f", width_pt=0.8)
    separator.paragraph_format.space_after = Pt(20)
    
    # === PREMIUM CONTENT AREA ===
    
    # Spacious content area with luxurious line spacing
    for i in range(14):
        body_para = doc.add_paragraph()
        body_para.paragraph_format.line_spacing = 1.4
        body_para.space_after = Pt(0)
    
    # === ELEGANT FOOTER SECTION ===
    
    # Pre-footer spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Dual geometric lines for premium footer effect
    footer_geo_1 = doc.add_paragraph()
    add_geometric_line(footer_geo_1, color_hex="#d4b37f", width_pt=1.5)
    footer_geo_1.paragraph_format.space_after = Pt(2)
    
    footer_geo_2 = doc.add_paragraph()
    add_geometric_line(footer_geo_2, color_hex="#e8dcc8", width_pt=0.8)
    footer_geo_2.paragraph_format.space_after = Pt(12)
    
    # Footer content - Line 1: Company & Locations
    footer_line1 = doc.add_paragraph()
    footer_line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_line1.paragraph_format.space_after = Pt(6)
    
    # Company name
    company_run = footer_line1.add_run('AHK')
    company_run.font.name = 'Montserrat'
    company_run.font.size = Pt(9.5)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(*hex_to_rgb('#111111'))
    
    strategies_run = footer_line1.add_run('Strategies')
    strategies_run.font.name = 'Montserrat'
    strategies_run.font.size = Pt(9.5)
    strategies_run.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    
    # Diamond separator
    separator_run = footer_line1.add_run('  ◆  ')
    separator_run.font.name = 'Segoe UI Symbol'
    separator_run.font.size = Pt(7)
    separator_run.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    # Location icon
    location_icon = footer_line1.add_run('📍 ')
    location_icon.font.name = 'Segoe UI Emoji'
    location_icon.font.size = Pt(8)
    
    # Locations
    location_run = footer_line1.add_run('Cairo  •  Dubai  •  Amman')
    location_run.font.name = 'Inter'
    location_run.font.size = Pt(8.5)
    location_run.font.color.rgb = RGBColor(*hex_to_rgb('#333333'))
    
    # Footer Line 2: Contact details
    footer_line2 = doc.add_paragraph()
    footer_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_line2.paragraph_format.space_after = Pt(8)
    
    # Phone icon
    phone_icon = footer_line2.add_run('☎ ')
    phone_icon.font.name = 'Segoe UI Symbol'
    phone_icon.font.size = Pt(8.5)
    phone_icon.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    # Phone number
    phone_run = footer_line2.add_run('+20 104 078 7571')
    phone_run.font.name = 'Inter'
    phone_run.font.size = Pt(8.5)
    phone_run.font.color.rgb = RGBColor(*hex_to_rgb('#333333'))
    
    # Separator
    sep2_run = footer_line2.add_run('  •  ')
    sep2_run.font.name = 'Inter'
    sep2_run.font.size = Pt(8.5)
    sep2_run.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    # Website icon
    web_icon = footer_line2.add_run('🌐 ')
    web_icon.font.name = 'Segoe UI Emoji'
    web_icon.font.size = Pt(8.5)
    
    # Website
    web_run = footer_line2.add_run('www.ahkstrategies.net')
    web_run.font.name = 'Inter'
    web_run.font.size = Pt(8.5)
    web_run.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    web_run.font.italic = True
    
    # Set premium default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(12)
    font.color.rgb = RGBColor(*hex_to_rgb('#1a1a1a'))
    
    # Heading styles
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Cormorant Garamond'
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(*hex_to_rgb('#111111'))
    
    # Save DOCX
    import datetime
    timestamp = datetime.datetime.now().strftime("%H%M%S")
    output_path_docx = rf"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\AHKStrategies_Letterhead_Master_v1.docx"
    
    try:
        doc.save(output_path_docx)
    except PermissionError:
        output_path_docx = rf"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\AHKStrategies_Letterhead_Master_FIXED_{timestamp}.docx"
        doc.save(output_path_docx)
        print("⚠️  Original file is open - saved as FIXED version")
    
    print("=" * 80)
    print("✨ MASTER LETTERHEAD CREATED ✨")
    print("=" * 80)
    print(f"📄 DOCX: {output_path_docx}")
    print(f"🎨 Typography: Cormorant Garamond + Montserrat + Garamond")
    print(f"⭐ Design: ULTIMATE SOPHISTICATION")
    print(f"🏆 Status: READY TO IMPRESS WORLD LEADERS")
    print("=" * 80)
    
    return output_path_docx

if __name__ == "__main__":
    create_master_letterhead()
