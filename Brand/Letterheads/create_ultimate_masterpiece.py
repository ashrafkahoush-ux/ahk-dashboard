"""
AHKStrategies Ultimate Letterhead - MASTERPIECE EDITION
Creating elegant vector graphics and premium design from SCRATCH
No gray backgrounds, no compromises, pure excellence
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import os

def create_elegant_banner():
    """Create a beautiful transparent brain-circuit banner from scratch"""
    
    # Create transparent image
    width, height = 2400, 400
    img = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    
    # Gold color
    gold = (212, 179, 127, 255)  # #d4b37f
    
    # Draw elegant horizontal lines (circuit paths)
    line_y = height // 2
    line_thickness = 4
    
    # Left side circuit lines
    for i in range(3):
        y_offset = (i - 1) * 30
        draw.rectangle([50, line_y + y_offset - 2, 800, line_y + y_offset + 2], fill=gold)
        # Add circuit nodes
        for x in range(100, 800, 120):
            draw.ellipse([x-6, line_y + y_offset - 6, x+6, line_y + y_offset + 6], fill=gold)
    
    # Right side circuit lines
    for i in range(3):
        y_offset = (i - 1) * 30
        draw.rectangle([1600, line_y + y_offset - 2, 2350, line_y + y_offset + 2], fill=gold)
        # Add circuit nodes
        for x in range(1650, 2350, 120):
            draw.ellipse([x-6, line_y + y_offset - 6, x+6, line_y + y_offset + 6], fill=gold)
    
    # Center brain symbol - HUMAN BRAIN SHAPE
    center_x, center_y = width // 2, height // 2
    
    # Draw stylized human brain (two hemispheres with curves)
    brain_width = 140
    brain_height = 100
    
    # Left hemisphere
    left_brain = [
        (center_x - 5, center_y - brain_height//2),  # Top center
        (center_x - brain_width//4, center_y - brain_height//2 - 10),  # Top curve out
        (center_x - brain_width//2, center_y - brain_height//3),  # Upper side
        (center_x - brain_width//2 - 10, center_y),  # Middle bulge
        (center_x - brain_width//2, center_y + brain_height//3),  # Lower side
        (center_x - brain_width//4, center_y + brain_height//2 + 10),  # Bottom curve
        (center_x - 5, center_y + brain_height//2),  # Bottom center
    ]
    
    # Right hemisphere
    right_brain = [
        (center_x + 5, center_y - brain_height//2),  # Top center
        (center_x + brain_width//4, center_y - brain_height//2 - 10),  # Top curve out
        (center_x + brain_width//2, center_y - brain_height//3),  # Upper side
        (center_x + brain_width//2 + 10, center_y),  # Middle bulge
        (center_x + brain_width//2, center_y + brain_height//3),  # Lower side
        (center_x + brain_width//4, center_y + brain_height//2 + 10),  # Bottom curve
        (center_x + 5, center_y + brain_height//2),  # Bottom center
    ]
    
    # Draw left hemisphere
    draw.line(left_brain, fill=gold, width=4, joint='curve')
    
    # Draw right hemisphere
    draw.line(right_brain, fill=gold, width=4, joint='curve')
    
    # Add brain folds (sulci) for realism
    # Left hemisphere folds
    draw.arc([center_x - brain_width//2 - 5, center_y - 30, center_x - 10, center_y + 10], 
             start=200, end=340, fill=gold, width=2)
    draw.arc([center_x - brain_width//2, center_y - 10, center_x - 15, center_y + 30], 
             start=180, end=320, fill=gold, width=2)
    
    # Right hemisphere folds
    draw.arc([center_x + 10, center_y - 30, center_x + brain_width//2 + 5, center_y + 10], 
             start=200, end=340, fill=gold, width=2)
    draw.arc([center_x + 15, center_y - 10, center_x + brain_width//2, center_y + 30], 
             start=220, end=360, fill=gold, width=2)
    
    # Center dividing line (corpus callosum)
    draw.line([center_x, center_y - brain_height//2, center_x, center_y + brain_height//2], 
              fill=gold, width=2)
    
    # Add small circuit connections to center brain
    draw.line([800, line_y, center_x - brain_width//2 - 20, center_y], fill=gold, width=3)
    draw.line([1600, line_y, center_x + brain_width//2 + 20, center_y], fill=gold, width=3)
    
    # Add decorative nodes
    for i in range(-3, 4):
        if i != 0:
            y = center_y + i * 40
            draw.ellipse([800 - 8, y - 8, 800 + 8, y + 8], fill=gold)
            draw.ellipse([1600 - 8, y - 8, 1600 + 8, y + 8], fill=gold)
    
    # Save transparent PNG
    output_path = r"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\banner_masterpiece.png"
    img.save(output_path, "PNG")
    print(f"✨ Masterpiece banner created: {output_path}")
    return output_path

def add_geometric_line(paragraph, color_hex="#d4b37f", width_pt=1.0):
    """Add elegant line"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(width_pt * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex.strip('#'))
    
    pBdr.append(bottom)
    pPr.append(pBdr)

def hex_to_rgb(hex_color):
    """Convert hex to RGB"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def create_ultimate_letterhead():
    """Create THE ULTIMATE letterhead - pure masterpiece"""
    
    print("=" * 80)
    print("🎨 CREATING MASTERPIECE FROM SCRATCH...")
    print("=" * 80)
    
    # Create banner first
    banner_path = create_elegant_banner()
    
    # Initialize document
    doc = Document()
    
    # Set page size and margins - A4
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    # === HEADER ===
    
    top_space = doc.add_paragraph()
    top_space.paragraph_format.space_after = Pt(6)
    
    # Company Name
    company_name = doc.add_paragraph()
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_name.paragraph_format.space_before = Pt(0)
    company_name.paragraph_format.space_after = Pt(4)
    
    name_run = company_name.add_run('A H K S T R A T E G I E S')
    name_run.font.name = 'Cormorant Garamond'
    name_run.font.size = Pt(15)
    name_run.font.color.rgb = RGBColor(*hex_to_rgb('#1a1a1a'))
    name_run.font.bold = False
    
    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_before = Pt(0)
    tagline.paragraph_format.space_after = Pt(8)
    
    tagline_run = tagline.add_run('Strategic   •   Human   •   Intelligent')
    tagline_run.font.name = 'Montserrat'
    tagline_run.font.size = Pt(8)
    tagline_run.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    
    # Masterpiece Banner
    banner_para = doc.add_paragraph()
    banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_para.paragraph_format.space_before = Pt(0)
    banner_para.paragraph_format.space_after = Pt(12)
    
    if os.path.exists(banner_path):
        run = banner_para.add_run()
        run.add_picture(banner_path, width=Cm(15))
    
    # Separator
    separator = doc.add_paragraph()
    add_geometric_line(separator, color_hex="#d4b37f", width_pt=1.0)
    separator.paragraph_format.space_after = Pt(22)
    
    # === CONTENT AREA ===
    for _ in range(15):
        body_para = doc.add_paragraph()
        body_para.paragraph_format.line_spacing = 1.4
        body_para.space_after = Pt(0)
    
    # === FOOTER ===
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Footer lines
    footer_line1 = doc.add_paragraph()
    add_geometric_line(footer_line1, color_hex="#d4b37f", width_pt=1.2)
    footer_line1.paragraph_format.space_after = Pt(2)
    
    footer_line2 = doc.add_paragraph()
    add_geometric_line(footer_line2, color_hex="#e8dcc8", width_pt=0.6)
    footer_line2.paragraph_format.space_after = Pt(10)
    
    # Footer content - Line 1
    footer1 = doc.add_paragraph()
    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer1.paragraph_format.space_after = Pt(5)
    
    f1_company = footer1.add_run('AHK')
    f1_company.font.name = 'Montserrat'
    f1_company.font.size = Pt(9)
    f1_company.font.bold = True
    f1_company.font.color.rgb = RGBColor(*hex_to_rgb('#111111'))
    
    f1_strategies = footer1.add_run('Strategies')
    f1_strategies.font.name = 'Montserrat'
    f1_strategies.font.size = Pt(9)
    f1_strategies.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    
    f1_sep = footer1.add_run('  ◆  ')
    f1_sep.font.name = 'Segoe UI Symbol'
    f1_sep.font.size = Pt(7)
    f1_sep.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    f1_loc_icon = footer1.add_run('📍 ')
    f1_loc_icon.font.size = Pt(8)
    
    f1_locations = footer1.add_run('Cairo  •  Dubai  •  Amman')
    f1_locations.font.name = 'Inter'
    f1_locations.font.size = Pt(8.5)
    f1_locations.font.color.rgb = RGBColor(*hex_to_rgb('#333333'))
    
    # Footer content - Line 2
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.paragraph_format.space_after = Pt(8)
    
    f2_phone_icon = footer2.add_run('☎ ')
    f2_phone_icon.font.name = 'Segoe UI Symbol'
    f2_phone_icon.font.size = Pt(8.5)
    f2_phone_icon.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    f2_phone = footer2.add_run('+20 104 078 7571')
    f2_phone.font.name = 'Inter'
    f2_phone.font.size = Pt(8.5)
    f2_phone.font.color.rgb = RGBColor(*hex_to_rgb('#333333'))
    
    f2_sep = footer2.add_run('  •  ')
    f2_sep.font.name = 'Inter'
    f2_sep.font.size = Pt(8.5)
    f2_sep.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    f2_web_icon = footer2.add_run('🌐 ')
    f2_web_icon.font.size = Pt(8.5)
    
    f2_web = footer2.add_run('www.ahkstrategies.net')
    f2_web.font.name = 'Inter'
    f2_web.font.size = Pt(8.5)
    f2_web.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    f2_web.font.italic = True
    
    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(11.5)
    font.color.rgb = RGBColor(*hex_to_rgb('#1a1a1a'))
    
    # Save
    output_path = r"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\AHKStrategies_Letterhead_MASTERPIECE.docx"
    doc.save(output_path)
    
    print("\n" + "=" * 80)
    print("✨✨✨ MASTERPIECE COMPLETE ✨✨✨")
    print("=" * 80)
    print(f"📄 File: {output_path}")
    print(f"🎨 Banner: Custom-created, transparent, elegant")
    print(f"⭐ Quality: UNCOMPROMISING EXCELLENCE")
    print(f"🏆 Status: MY OWN MASTERPIECE - READY TO IMPRESS")
    print("=" * 80)
    
    return output_path

if __name__ == "__main__":
    create_ultimate_letterhead()
