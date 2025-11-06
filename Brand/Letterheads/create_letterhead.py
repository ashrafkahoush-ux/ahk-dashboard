"""
AHKStrategies Letterhead Generator - MASTERPIECE EDITION
Creates a stunning, award-worthy corporate letterhead with sophisticated design elements
Premium geometric patterns, elegant spacing, and visual excellence
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OxmlElementShared
import os

def add_geometric_line(paragraph, color_hex="#d4b37f", width_pt=1.5, style='single'):
    """Add a sophisticated line to a paragraph using border"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), style)
    bottom.set(qn('w:sz'), str(int(width_pt * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex.strip('#'))
    
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_top_line(paragraph, color_hex="#d4b37f", width_pt=1.5):
    """Add a line to the top of a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(int(width_pt * 8)))
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), color_hex.strip('#'))
    
    pBdr.append(top)
    pPr.append(pBdr)

def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def add_shading(paragraph, color_hex="#f9f7f4"):
    """Add subtle background shading to a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex.strip('#'))
    pPr.append(shd)

def create_letterhead():
    """Create the AHKStrategies MASTERPIECE letterhead document"""
    
    # Initialize document
    doc = Document()
    
    # Set page margins - A4 size with optimized margins for premium feel
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)  # A4 height
        section.page_width = Cm(21)     # A4 width
        section.top_margin = Cm(2.2)    # More breathing room
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)   # Premium left margin
        section.right_margin = Cm(2.0)
    
    # === SOPHISTICATED HEADER SECTION ===
    
    # Top accent bar - Dual-tone geometric design
    top_accent = doc.add_paragraph()
    top_accent.paragraph_format.space_before = Pt(0)
    top_accent.paragraph_format.space_after = Pt(0)
    add_geometric_line(top_accent, color_hex="#d4b37f", width_pt=2.5)
    
    # Subtle spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Logo and Brand Identity Section
    logo_section = doc.add_table(rows=1, cols=2)
    logo_section.autofit = False
    logo_section.allow_autofit = False
    
    # Set column widths for perfect balance
    logo_section.columns[0].width = Cm(9)
    logo_section.columns[1].width = Cm(8)
    
    # Left cell - Logo with premium spacing
    left_cell = logo_section.rows[0].cells[0]
    left_cell.vertical_alignment = 1  # Center
    logo_paragraph = left_cell.paragraphs[0]
    logo_paragraph.paragraph_format.space_before = Pt(12)
    logo_paragraph.paragraph_format.space_after = Pt(12)
    
    # Add logo with optimal size
    logo_path = r"c:\Users\ashra\OneDrive\Desktop\AHK Profile\AHK-Digital-ID-Kit\ahkstrategies-logo.png"
    if os.path.exists(logo_path):
        run = logo_paragraph.add_run()
        run.add_picture(logo_path, width=Cm(4.2))  # Larger, more prominent
    
    # Right cell - Sophisticated geometric pattern
    right_cell = logo_section.rows[0].cells[1]
    right_cell.vertical_alignment = 1  # Center
    
    # Multi-line geometric accent
    for i in range(3):
        geo_para = right_cell.paragraphs[0] if i == 0 else right_cell.add_paragraph()
        geo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        geo_para.paragraph_format.space_after = Pt(2)
        
        # Create layered geometric lines with different opacities
        geo_chars = ['━━━', '━━', '━'][i]
        geo_run = geo_para.add_run(geo_chars)
        geo_run.font.size = Pt(16 - i*2)
        
        # Gradient effect through opacity simulation
        opacity_colors = ['#d4b37f', '#dcc399', '#e6d4b8'][i]
        geo_run.font.color.rgb = RGBColor(*hex_to_rgb(opacity_colors))
        geo_run.font.name = 'Calibri'
    
    # Remove table borders for clean look
    for row in logo_section.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    # Refined spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Premium tagline with visual sophistication
    tagline_container = doc.add_paragraph()
    tagline_container.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tagline_container.paragraph_format.left_indent = Cm(0.15)
    tagline_container.paragraph_format.space_before = Pt(4)
    tagline_container.paragraph_format.space_after = Pt(18)
    
    # Tagline text with letter spacing effect
    tagline_run = tagline_container.add_run('S T R A T E G I C   •   H U M A N   •   I N T E L L I G E N T')
    tagline_run.font.name = 'Montserrat'
    tagline_run.font.size = Pt(8.5)
    tagline_run.font.bold = True
    tagline_run.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))  # Muted gold for sophistication
    
    # Elegant divider line after header
    divider = doc.add_paragraph()
    add_geometric_line(divider, color_hex="#e8dcc8", width_pt=0.8)
    divider.paragraph_format.space_after = Pt(28)
    
    
    # === PREMIUM CONTENT AREA ===
    
    # Main content space with sophisticated spacing
    for i in range(12):
        body_para = doc.add_paragraph()
        body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        body_para.paragraph_format.line_spacing = 1.35  # Luxurious line spacing
        body_para.space_after = Pt(0)
    
    # === EXQUISITE FOOTER SECTION ===
    
    # Pre-footer spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Sophisticated footer geometric design - Multi-layer
    footer_geo_1 = doc.add_paragraph()
    add_geometric_line(footer_geo_1, color_hex="#d4b37f", width_pt=2.0)
    footer_geo_1.paragraph_format.space_after = Pt(2)
    
    footer_geo_2 = doc.add_paragraph()
    add_geometric_line(footer_geo_2, color_hex="#e8dcc8", width_pt=1.0)
    footer_geo_2.paragraph_format.space_after = Pt(14)
    
    # Premium footer content - Company name and locations
    footer_line1 = doc.add_paragraph()
    footer_line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_line1.paragraph_format.space_after = Pt(6)
    
    # Company name in bold
    company_run = footer_line1.add_run('AHK')
    company_run.font.name = 'Montserrat'
    company_run.font.size = Pt(9.5)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(*hex_to_rgb('#111111'))
    
    strategies_run = footer_line1.add_run('Strategies')
    strategies_run.font.name = 'Montserrat'
    strategies_run.font.size = Pt(9.5)
    strategies_run.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    
    # Separator with style
    separator_run = footer_line1.add_run('  ◆  ')
    separator_run.font.name = 'Segoe UI Symbol'
    separator_run.font.size = Pt(7)
    separator_run.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    # Location icon (elegant pin)
    location_icon = footer_line1.add_run('📍 ')
    location_icon.font.name = 'Segoe UI Emoji'
    location_icon.font.size = Pt(8)
    
    # Locations
    location_run = footer_line1.add_run('Cairo  •  Dubai  •  Amman')
    location_run.font.name = 'Inter'
    location_run.font.size = Pt(8.5)
    location_run.font.color.rgb = RGBColor(*hex_to_rgb('#333333'))
    
    # Second line - Contact details with elegant icons
    footer_line2 = doc.add_paragraph()
    footer_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_line2.paragraph_format.space_after = Pt(10)
    
    # Phone icon (sophisticated)
    phone_icon = footer_line2.add_run('☎ ')
    phone_icon.font.name = 'Segoe UI Symbol'
    phone_icon.font.size = Pt(8.5)
    phone_icon.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    # Phone number
    phone_run = footer_line2.add_run('+20 104 078 7571')
    phone_run.font.name = 'Inter'
    phone_run.font.size = Pt(8.5)
    phone_run.font.color.rgb = RGBColor(*hex_to_rgb('#333333'))
    
    # Elegant separator
    sep2_run = footer_line2.add_run('  •  ')
    sep2_run.font.name = 'Inter'
    sep2_run.font.size = Pt(8.5)
    sep2_run.font.color.rgb = RGBColor(*hex_to_rgb('#d4b37f'))
    
    # Website icon (globe)
    web_icon = footer_line2.add_run('🌐 ')
    web_icon.font.name = 'Segoe UI Emoji'
    web_icon.font.size = Pt(8.5)
    
    # Website
    web_run = footer_line2.add_run('www.ahkstrategies.net')
    web_run.font.name = 'Inter'
    web_run.font.size = Pt(8.5)
    web_run.font.color.rgb = RGBColor(*hex_to_rgb('#8b7355'))
    web_run.font.italic = True
    
    # Bottom accent - Mirror of top
    bottom_accent = doc.add_paragraph()
    bottom_accent.paragraph_format.space_before = Pt(4)
    bottom_accent.paragraph_format.space_after = Pt(0)
    add_geometric_line(bottom_accent, color_hex="#d4b37f", width_pt=1.5)
    
    # Set premium default font styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11.5)  # Slightly larger for readability
    font.color.rgb = RGBColor(*hex_to_rgb('#1a1a1a'))  # Rich black
    
    # Heading styles for professional documents
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Montserrat'
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(*hex_to_rgb('#111111'))
    
    # Save the MASTERPIECE
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = rf"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\AHKStrategies_Letterhead_v1_{timestamp}.docx"
    doc.save(output_path)
    
    # Also try to save the main version (if not locked)
    main_path = r"C:\Users\ashra\OneDrive\Desktop\AHK_Dashboard_v1\Brand\Letterheads\AHKStrategies_Letterhead_v1.docx"
    try:
        doc.save(main_path)
        output_path = main_path
    except PermissionError:
        print("⚠️  Original file is open - saved timestamped version instead")
    
    print("=" * 70)
    print("✨ MASTERPIECE CREATED ✨")
    print("=" * 70)
    print(f"📄 File: {output_path}")
    print(f"🎨 Design Level: PREMIUM EXCELLENCE")
    print(f"⭐ Status: READY TO WOW")
    print("=" * 70)
    
    return output_path

if __name__ == "__main__":
    create_letterhead()
